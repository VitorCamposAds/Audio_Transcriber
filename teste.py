import sys
import os
import subprocess
import time
import whisper
from docx import Document
from PyQt5.QtWidgets import (
    QApplication,
    QMainWindow,
    QWidget,
    QLabel,
    QPushButton,
    QFileDialog,
    QMessageBox,
    QVBoxLayout,
    QProgressBar
)
from PyQt5.QtGui import QColor, QPalette
from PyQt5.QtCore import Qt, QStandardPaths

# Função para extrair o áudio do arquivo de vídeo usando FFmpeg
def extract_audio(input_video_file, output_audio_file):
    try:
        subprocess.run(['ffmpeg', '-i', input_video_file, '-vn', '-acodec', 'copy', output_audio_file], check=True)
        QMessageBox.information(None, "Extração de Áudio", "Áudio extraído com sucesso. Aguarde a transcrição!")
    except subprocess.CalledProcessError as e:
        QMessageBox.critical(None, "Erro", f"Erro ao extrair áudio do vídeo: {str(e)}")

# Função para transcrever o áudio em texto e salvar em DOCX
def transcribe_and_save(input_audio_file, progress_bar):
    try:
        modelo = whisper.load_model("small")
        
        # Inicializar a barra de progresso
        progress_bar.setValue(0)
        
        # Iniciar contagem de tempo para sincronizar a barra de progresso
        start_time = time.time()
        
        # Verificar o caminho do arquivo de áudio antes de prosseguir
        print(f"Caminho do arquivo de áudio: {input_audio_file}")
        
        # Verificar se o arquivo de áudio existe
        if not os.path.exists(input_audio_file):
            raise FileNotFoundError("Arquivo de áudio não encontrado.")
        
        # Transcrição do áudio
        resposta = modelo.transcribe(input_audio_file, fp16=False)
        
        # Salvar a transcrição em DOCX no desktop do usuário
        desktop_path = QStandardPaths.standardLocations(QStandardPaths.DesktopLocation)[0]
        output_docx_file = os.path.join(desktop_path, "Transcrição.docx")
        
        doc = Document()
        doc.add_heading('Transcrição de Áudio', level=1)
        doc.add_paragraph(resposta["text"])
        doc.save(output_docx_file)
        
        # Calcular o tempo total de transcrição
        end_time = time.time()
        total_time = end_time - start_time
        
        QMessageBox.information(None, "Transcrição Concluída", f"Transcrição salva em: {output_docx_file}.\nTempo total de transcrição: {total_time:.2f} segundos")
        progress_bar.setValue(100)  # Define o valor para 100% após a conclusão
        
    except FileNotFoundError:
        QMessageBox.critical(None, "Erro", "Arquivo de áudio não encontrado.")
    except whisper.AudioFormatError:
        QMessageBox.critical(None, "Erro", "Formato de áudio não suportado.")
    except Exception as e:
        QMessageBox.critical(None, "Erro inesperado", f"Erro inesperado: {e}")

# Classe principal da aplicação PyQt
class AudioTranscriberApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle("Extração de Áudio e Transcrição para DOCX")
        self.setGeometry(100, 100, 600, 300)  # Ajuste das dimensões da janela
        
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        vbox = QVBoxLayout()

        # Rótulo explicativo
        label_intro = QLabel("Selecione um arquivo de vídeo para extrair áudio e transcrever para DOCX:")
        vbox.addWidget(label_intro)
        
        # Botão para extrair áudio do vídeo, transcrever e salvar em DOCX
        extract_button = QPushButton("Selecionar Vídeo e Extrair Áudio")
        extract_button.clicked.connect(self.handle_extract_and_transcribe)
        vbox.addWidget(extract_button)
        
        # Espaço vertical para layout
        vbox.addStretch(1)
        
        # Rótulo para indicar o progresso
        self.progress_label = QLabel("")
        vbox.addWidget(self.progress_label)
        
        # Barra de progresso
        self.progress_bar = QProgressBar()
        self.progress_bar.setTextVisible(False)  # Ocultar texto dentro da barra de progresso
        
        # Configuração da cor da barra de progresso
        palette = self.progress_bar.palette()
        palette.setColor(QPalette.Highlight, QColor(42, 130, 218))  # Cor azul para o destaque
        self.progress_bar.setPalette(palette)
        
        vbox.addWidget(self.progress_bar)

        central_widget.setLayout(vbox)

    def handle_extract_and_transcribe(self):
        input_video_file, _ = QFileDialog.getOpenFileName(self, "Selecionar Vídeo", "", "Arquivos de Vídeo (*.mp4)")
        if input_video_file:
            output_dir = QStandardPaths.standardLocations(QStandardPaths.DesktopLocation)[0]
            
            # Nome e localização do arquivo de áudio extraído
            audio_base_name = os.path.splitext(os.path.basename(input_video_file))[0]
            audio_temp_file = os.path.join(output_dir, f"{audio_base_name}.aac")
            
            # Extrair áudio do vídeo
            extract_audio(input_video_file, audio_temp_file)
            
            # Transcrever áudio e salvar em DOCX
            self.progress_label.setText("Processando transcrição...")
            self.progress_bar.setRange(0, 0)  # Indeterminado enquanto processa
            transcribe_and_save(audio_temp_file, self.progress_bar)
            self.progress_label.setText("")
            self.progress_bar.setRange(0, 100)  # Restaura a barra de progresso

# Função principal para iniciar o aplicativo
def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")  # Aplicar o estilo Fusion para uma aparência mais moderna

    # Definir o estilo do aplicativo para um esquema de cores personalizado
    dark_palette = QPalette()
    dark_palette.setColor(QPalette.Window, QColor(53, 53, 53))
    dark_palette.setColor(QPalette.WindowText, Qt.white)
    dark_palette.setColor(QPalette.Base, QColor(25, 25, 25))
    dark_palette.setColor(QPalette.AlternateBase, QColor(53, 53, 53))
    dark_palette.setColor(QPalette.ToolTipBase, Qt.white)
    dark_palette.setColor(QPalette.ToolTipText, Qt.white)
    dark_palette.setColor(QPalette.Text, Qt.white)
    dark_palette.setColor(QPalette.Button, QColor(53, 53, 53))
    dark_palette.setColor(QPalette.ButtonText, Qt.white)
    dark_palette.setColor(QPalette.BrightText, Qt.red)
    dark_palette.setColor(QPalette.Link, QColor(42, 130, 218))
    dark_palette.setColor(QPalette.Highlight, QColor(42, 130, 218))
    dark_palette.setColor(QPalette.HighlightedText, Qt.black)
    app.setPalette(dark_palette)

    window = AudioTranscriberApp()
    window.show()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()